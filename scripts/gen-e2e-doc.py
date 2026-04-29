"""
產生 Sprint 5 E2E 驗證教學 Word 文件
輸出：docs/sprint5-e2e-verification.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = os.path.join(os.path.dirname(__file__), '..', 'docs', 'sprint5-e2e-verification.docx')

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(table, border_color='D0D0D0'):
    """Light border for all cells."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_body(doc, text, bold=False, color=None, left_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    return p

def add_code(doc, text, left_indent=1.0):
    """Monospace code block (light gray bg via paragraph shading)."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(left_indent)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        # Gray background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)

def add_checklist(doc, items, left_indent=1.0):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(left_indent)
        run = p.add_run('☐  ' + item)
        run.font.size = Pt(11)

def add_table(doc, headers, rows, header_bg='2D5DA1', header_fg='FFFFFF'):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(*bytes.fromhex(header_fg))
                run.font.size = Pt(10.5)

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'F7F9FC'
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            cell.text = cell_text
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)

    set_cell_border(table)
    return table

def add_step_header(doc, number: str, title: str):
    p = doc.add_paragraph()
    r1 = p.add_run(f'Step {number}  ')
    r1.font.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x2D, 0x5D, 0xA1)
    r2 = p.add_run(title)
    r2.font.size = Pt(12)
    r2.font.bold = False
    return p

def add_note(doc, text, icon='ℹ️'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'{icon}  {text}')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_pass_box(doc, items):
    """Green-tinted pass criteria box."""
    p = doc.add_paragraph()
    r = p.add_run('✅  通過條件')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x7A, 0x3C)
    for item in items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Cm(1)
        run = bp.add_run('☐  ' + item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x3C)

# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles['Normal'].font.name = '微軟正黑體'
doc.styles['Normal'].font.size = Pt(11)

# ── Cover ─────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sprint 5 — E2E 驗證教學')
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x2D, 0x5D, 0xA1)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('跨 Agent 通訊機制（SendMessage MCP Server）\n手動驗收操作手冊')
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

meta_rows = [
    ('專案', 'AgentHub'),
    ('Sprint', 'Sprint 5'),
    ('文件類型', 'E2E 人工驗收手冊（T8 / T9 / T10）'),
    ('日期', '2026-04-29'),
    ('撰寫人', 'tech-lead'),
]
t = add_table(doc, ['欄位', '值'], meta_rows, header_bg='2D5DA1')

doc.add_paragraph()

# ── Section 1: 前置準備 ───────────────────────────────────────────────────────

add_heading(doc, '一、前置準備', level=1, color='2D5DA1')
add_body(doc, '開始驗證前，請確認以下兩個條件都已滿足。', bold=False)
doc.add_paragraph()

add_heading(doc, '1-1  編譯 MCP Server', level=2)
add_body(doc, '在 AgentHub 專案根目錄執行：')
add_code(doc, '''cd "C:/Users/Bandai/Desktop/ALL PROJECT/Agent-hub"
npm run build:mcp''')
doc.add_paragraph()
add_body(doc, '確認編譯產物存在：')
add_code(doc, 'ls out/mcp/send-message-server.js')
add_note(doc, '若檔案不存在，session spawn 時 MCP server 不會注入，SendMessage 工具不可用。')
doc.add_paragraph()

add_heading(doc, '1-2  準備兩個視窗', level=2)
add_table(doc, ['視窗', '用途'], [
    ('AgentHub GUI', '開 session、觀察 PTY 輸出和 Tool call 結果'),
    ('Terminal / 檔案總管', '即時查看 inbox JSON 檔案'),
])
doc.add_paragraph()
add_body(doc, 'Inbox 存放位置：')
add_code(doc, 'C:\\Users\\Bandai\\.claude\\teams\\default\\inboxes\\')
doc.add_paragraph()

add_heading(doc, '1-3  快速查看 Inbox 指令', level=2)
add_body(doc, '在 Terminal 執行以下指令，可一次列出所有 agent 的 inbox 狀態：')
add_code(doc, r'''python "C:/Users/Bandai/anaconda3/python.exe" -c "
import json, os
inbox_dir = r'C:\Users\Bandai\.claude\teams\default\inboxes'
if not os.path.exists(inbox_dir):
    print('inbox 目錄不存在')
else:
    files = [f for f in os.listdir(inbox_dir) if f.endswith('.json')]
    if not files:
        print('inbox 目錄為空')
    for f in sorted(files):
        data = json.load(open(os.path.join(inbox_dir, f), encoding='utf-8'))
        print(f'{f}: {len(data)} 筆')
        if data:
            last = data[-1]
            print(f'  最新: from={last[\"from\"]}  text={last[\"text\"][:50]}')
"''')

doc.add_page_break()

# ── Section 2: T8 ────────────────────────────────────────────────────────────

add_heading(doc, '二、T8 — 4 層接力通訊驗證', level=1, color='2D5DA1')
add_body(doc, '目標：boss → product-manager → tech-lead → backend-architect，全鏈訊息傳遞成功。')
doc.add_paragraph()

add_table(doc, ['鏈路', 'Agent', '動作'], [
    ('Layer 1', 'boss', '呼叫 SendMessage → product-manager'),
    ('Layer 2', 'product-manager', '收到訊息 → 呼叫 SendMessage → tech-lead'),
    ('Layer 3', 'tech-lead', '收到訊息 → 呼叫 SendMessage → backend-architect'),
    ('Layer 4', 'backend-architect', '收到訊息（終點）'),
])
doc.add_paragraph()

# Steps
add_step_header(doc, '1', '開 boss session')
add_body(doc, 'AgentHub → 開新 session，Agent 選 boss。\n在對話框輸入：')
add_code(doc, '請用 SendMessage 工具寄一封測試訊息給 product-manager，\n內容：「T8 E2E 測試：請你轉寄給 tech-lead」')
add_note(doc, '等待 agent 執行 send_message tool call，畫面應顯示 tool use 區塊。')
doc.add_paragraph()

add_step_header(doc, '2', '驗證 product-manager inbox（等 3 秒後）')
add_code(doc, r'cat "C:\Users\Bandai\.claude\teams\default\inboxes\product-manager.json"')
add_body(doc, '預期看到：')
add_code(doc, '''{
  "from": "boss",
  "text": "T8 E2E 測試：請你轉寄給 tech-lead",
  "read": false,
  "messageId": "...",
  "timestamp": "..."
}''')
doc.add_paragraph()

add_step_header(doc, '3', '開 product-manager session')
add_body(doc, '開新 session，Agent 選 product-manager。輸入：')
add_code(doc, '請用 list_inbox 查看你的收件匣，\n然後把收到的訊息轉寄給 tech-lead\n（SendMessage 工具，內容：「T8 接力：來自 boss，請轉給 backend-architect」）')
doc.add_paragraph()

add_step_header(doc, '4', '驗證 tech-lead inbox')
add_code(doc, r'cat "C:\Users\Bandai\.claude\teams\default\inboxes\tech-lead.json"')
doc.add_paragraph()

add_step_header(doc, '5', '開 tech-lead session → 轉寄 backend-architect')
add_body(doc, '開新 session，Agent 選 tech-lead。輸入：')
add_code(doc, '請用 list_inbox 查看收件匣，\n並把訊息用 SendMessage 轉寄給 backend-architect\n（內容：「T8 接力完成通知」）')
doc.add_paragraph()

add_step_header(doc, '6', '驗證 backend-architect inbox（最終確認）')
add_code(doc, r'cat "C:\Users\Bandai\.claude\teams\default\inboxes\backend-architect.json"')
doc.add_paragraph()

add_pass_box(doc, [
    'product-manager.json 有 from: "boss" 的訊息',
    'tech-lead.json 有 from: "product-manager" 的訊息',
    'backend-architect.json 有 from: "tech-lead" 的訊息',
    '每筆訊息均包含：from / text / summary / timestamp / read / messageId',
    '訊息 read 欄位為 false',
    '每層訊息延遲 ≤ 5 秒（含 3 秒 inbox poll interval）',
])

doc.add_page_break()

# ── Section 3: T9 ────────────────────────────────────────────────────────────

add_heading(doc, '三、T9 — 越權攔截驗證', level=1, color='2D5DA1')
add_body(doc, '目標：academic 部門 L2 agent 嘗試寄訊息給 engineering 部門 L2 agent → 回傳 UNAUTHORIZED，inbox 不變。')
doc.add_paragraph()

add_note(doc, '設計原則：agent 只能寄訊息給 manages（下屬）、reportsTo（上級）、coordinatesWith（協作）名單內的對象。跨部門 L2 不在任一清單，應被攔截。', icon='🔒')
doc.add_paragraph()

add_step_header(doc, '1', '開 literature-scout session（academic 部門 L2）')
add_body(doc, 'AgentHub → 開新 session，Agent 選 literature-scout（或 data-analyst 等 academic L2）。')
doc.add_paragraph()

add_step_header(doc, '2', '嘗試越權發送訊息')
add_body(doc, '在對話框輸入：')
add_code(doc, '請用 SendMessage 工具寄訊息給 backend-architect，\n內容：「越權測試訊息」')
doc.add_paragraph()

add_step_header(doc, '3', '觀察 Tool Result')
add_body(doc, 'agent 的 tool call 回傳應顯示：')
add_code(doc, '''{
  "error": "UNAUTHORIZED",
  "message": "Agent \\"literature-scout\\" is not allowed to send messages to \\"backend-architect\\". Allowed targets: [...]"
}''')
doc.add_paragraph()

add_step_header(doc, '4', '確認 inbox 未被寫入')
add_code(doc, r'cat "C:\Users\Bandai\.claude\teams\default\inboxes\backend-architect.json"')
add_body(doc, '預期：檔案不存在，或存在但無新增記錄。')
doc.add_paragraph()

add_pass_box(doc, [
    'tool result 顯示 "error": "UNAUTHORIZED"',
    'error message 包含越權對象名稱（backend-architect）',
    'backend-architect.json 無新增記錄（inbox 不變）',
])

doc.add_page_break()

# ── Section 4: T10 ───────────────────────────────────────────────────────────

add_heading(doc, '四、T10 — Rate Limit 驗證', level=1, color='2D5DA1')
add_body(doc, '目標：單一 session 連發 21 條訊息 → 前 20 條成功，第 21 條回傳 RATE_LIMITED。')
doc.add_paragraph()

add_note(doc, 'Rate limit 是 per-session、per-agentId、in-memory 追蹤，重啟 session 後計數重置。', icon='⏱️')
doc.add_paragraph()

add_step_header(doc, '1', '開 tech-lead session（有 allowedTargets 的 agent）')
add_body(doc, 'AgentHub → 開新 session，Agent 選 tech-lead。')
doc.add_paragraph()

add_step_header(doc, '2', '讓 agent 連發 21 條訊息')
add_body(doc, '在對話框輸入：')
add_code(doc, '請連續用 SendMessage 工具寄 21 條訊息給 product-manager，\n每條內容分別是「Rate limit test 1」到「Rate limit test 21」。\n每條都立刻寄，不要停頓。所有 tool call 完成後，告訴我第 21 條的結果。')
doc.add_paragraph()

add_step_header(doc, '3', '觀察 Tool Result 差異')
add_table(doc, ['訊息編號', '預期 tool result'], [
    ('第 1～20 條', '{ "status": "pending", "to": "product-manager", "message_id": "..." }'),
    ('第 21 條', '{ "error": "RATE_LIMITED", "message": "Rate limit exceeded: max 20 messages per hour per session." }'),
])
doc.add_paragraph()

add_step_header(doc, '4', '確認 inbox 精確有 20 筆')
add_code(doc, r'''"C:/Users/Bandai/anaconda3/python.exe" -c "
import json
data = json.load(open(r'C:\Users\Bandai\.claude\teams\default\inboxes\product-manager.json', encoding='utf-8'))
print(f'inbox 訊息數: {len(data)}')
print(f'最後一筆: {data[-1][\"text\"]}')
"''')
add_body(doc, '預期輸出：')
add_code(doc, 'inbox 訊息數: 20\n最後一筆: Rate limit test 20')
doc.add_paragraph()

add_pass_box(doc, [
    '前 20 條 tool result 均顯示 "status": "pending"',
    '第 21 條 tool result 顯示 "error": "RATE_LIMITED"',
    'product-manager.json 精確有 20 筆記錄（不多不少）',
])

doc.add_page_break()

# ── Section 5: 結果回報格式 ───────────────────────────────────────────────────

add_heading(doc, '五、驗收結果回報格式', level=1, color='2D5DA1')
add_body(doc, '每個 Test Case 驗證完成後，請提供以下資訊給 tech-lead 記錄到 dev-plan §10：')
doc.add_paragraph()
add_table(doc, ['Test Case', '結果', '截圖 / 附件'], [
    ('T8 4層接力', '通過 / 未通過', 'inbox JSON 截圖 × 3'),
    ('T9 越權攔截', '通過 / 未通過', 'tool result 截圖'),
    ('T10 Rate Limit', '通過 / 未通過', 'tool result 截圖 + inbox 筆數截圖'),
])
doc.add_paragraph()
add_note(doc, '三個 TC 全部通過後，tech-lead 將提交 G3 Gate → PM Review → 老闆審核。', icon='📋')

# ── Save ──────────────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f'[OK] Output: {os.path.abspath(OUT_PATH)}')
